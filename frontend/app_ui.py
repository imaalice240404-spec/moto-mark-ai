import streamlit as st
import sys
import os
import io
import pandas as pd
from docx import Document

# 連結後端大腦
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))
from backend.app import generate_risk_report

# --- 1. 頁面配置 ---
st.set_page_config(
    page_title="MotoMark AI | 企業智財中心",
    page_icon="🏍️",
    layout="wide"
)

# --- 2. CSS 視覺魔法 ---
st.markdown("""
    <style>
    .stApp { background-color: #F8FAFC; }
    div[data-testid="stVerticalBlock"] > div[style*="border"] {
        background-color: #FFFFFF;
        border-radius: 12px;
        border: 1px solid #E2E8F0;
        box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.05);
        padding: 24px;
    }
    .stButton>button[kind="primary"] {
        background-color: #0F172A;
        color: white;
        font-weight: 600;
        border-radius: 8px;
        transition: all 0.2s ease;
    }
    .stButton>button[kind="primary"]:hover {
        background-color: #334155;
        transform: translateY(-2px);
    }
    /* 讓 Tabs 標籤看起來更大更清楚 */
    .stTabs [data-baseweb="tab-list"] {
        gap: 24px;
    }
    .stTabs [data-baseweb="tab"] {
        height: 50px;
        white-space: pre-wrap;
        font-size: 18px;
        font-weight: 600;
    }
    header {visibility: hidden;}
    </style>
""", unsafe_allow_html=True)

# --- 3. Word 產生器 ---
def create_word_doc(report_text):
    doc = Document()
    doc.add_heading('MotoMark AI - 商標核駁風險評估意見書', 0)
    doc.add_paragraph(report_text.replace('#', '').replace('*', ''))
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- 4. 讀取真實 Excel 資料庫的函數 ---
@st.cache_data
def load_real_data():
    current_dir = os.path.dirname(os.path.abspath(__file__))
    file_path = os.path.join(current_dir, '..', 'data', 'trademarks.xlsx')
    try:
        all_sheets = pd.read_excel(file_path, sheet_name=None)
        df = pd.concat(all_sheets.values(), ignore_index=True)
        df = df.fillna("") # 把空白欄位填滿避免報錯
        return df
    except Exception as e:
        return pd.DataFrame() # 如果讀不到就回傳空表格

df_real = load_real_data()

# ==========================================
# 🌟 使用頂部 Tabs 取代側邊欄，畫面更直覺 🌟
# ==========================================
tab1, tab2, tab3 = st.tabs(["🏠 首頁：競爭對手情報", "⚡ 智財風險檢索中心", "📂 我的商標管理庫"])

# ------------------------------------------
# 頁面 1：首頁 (串接真實資料！)
# ------------------------------------------
with tab1:
    st.title("📊 競爭對手情報儀表板")
    st.markdown("本頁面資料已**即時連線至您的本機資料庫 (`trademarks.xlsx`)**，為您自動統整競爭對手申請動態。")
    
    if not df_real.empty:
        # 動態計算真實數據
        total_count = len(df_real)
        kymco_count = df_real['申請人'].astype(str).str.contains('光陽', na=False).sum()
        sym_count = df_real['申請人'].astype(str).str.contains('三陽', na=False).sum()
        yamaha_count = df_real['申請人'].astype(str).str.contains('山葉', na=False).sum()

        col1, col2, col3, col4 = st.columns(4)
        col1.metric("本地資料庫總案量", f"{total_count} 件")
        col2.metric("光陽 (KYMCO) 總案量", f"{kymco_count} 件")
        col3.metric("三陽 (SYM) 總案量", f"{sym_count} 件")
        col4.metric("台灣山葉 (YAMAHA)", f"{yamaha_count} 件")
        
        st.divider()
        st.subheader("🚨 真實資料庫預覽")
        st.markdown("以下為從您上傳的資料庫中讀取的真實資料內容：")
        
        with st.container(border=True):
            # 自動抓取我們需要的欄位來顯示，避免欄位太多看起來很亂
            cols_to_show = [c for c in ['商標名稱', '申請人', '申請日期', '類別', '申請案號', '商標文字'] if c in df_real.columns]
            if cols_to_show:
                st.dataframe(df_real[cols_to_show], use_container_width=True, height=400)
            else:
                st.dataframe(df_real, use_container_width=True, height=400)
    else:
        st.error("⚠️ 無法讀取 `data/trademarks.xlsx`，請確認檔案是否存在。")

# ------------------------------------------
# 頁面 2：智財風險檢索中心
# ------------------------------------------
with tab2:
    st.title("⚡ 智財風險檢索中心")
    st.markdown("輸入擬申請的商標資訊，系統將自動調閱 **RAG 歷史案卷** 並比對 **TIPO 前案資料庫**。 *(註：您的檢索紀錄全程加密，不會用於 AI 訓練)*")

    with st.container(border=True):
        col1, col2, col3 = st.columns([2, 2, 1])
        with col1:
            trademark_input = st.text_input("🔍 擬申請商標名稱", "NNBCU")
        with col2:
            product_class = st.selectbox(
                "📁 指定商品/服務類別", 
                [
                    "第 12 類 - 機車 (整車)",
                    "第 12+35 類 - 整車 ＋ 零售批發 (強防禦)",
                    "第 12+37 類 - 整車 ＋ 維修保養 (強防禦)",
                    "第 12+35+37 類 - 品牌全方位防護"
                ]
            )
        with col3:
            st.write("")
            st.write("")
            submit_btn = st.button("🚀 啟動深度分析", type="primary", use_container_width=True)

    if submit_btn:
        with st.spinner('🔄 系統封閉運算中：比對外觀讀音 ➔ 檢索本地歷史案卷 ➔ 彙整法理邏輯...'):
            report = generate_risk_report(trademark_input, product_class)
            st.session_state["last_report"] = report

    if "last_report" in st.session_state:
        st.divider()
        head_col1, head_col2 = st.columns([4, 1])
        with head_col1:
            st.subheader(f"📑 評估報告")
        with head_col2:
            word_file = create_word_doc(st.session_state["last_report"])
            st.download_button("📥 下載 Word 報告", data=word_file, file_name="風險評估書.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
        
        with st.container(border=True):
            st.markdown(st.session_state["last_report"])

# ------------------------------------------
# 頁面 3：我的商標管理庫
# ------------------------------------------
with tab3:
    st.title("📂 企業商標資產管理")
    st.markdown("在此輸入您的公司名稱，系統將自動從資料庫中撈出屬於您的商標資產進行集中管理。")
    
    with st.container(border=True):
        my_company = st.text_input("🏢 輸入您的申請人名稱 (例如輸入「光陽」或「三陽」來測試)：", "光陽")
        
        if my_company and not df_real.empty:
            # 使用 Pandas 語法：篩選出申請人包含關鍵字的資料
            my_df = df_real[df_real['申請人'].astype(str).str.contains(my_company, na=False)]
            
            st.success(f"✅ 找到 {len(my_df)} 筆屬於「{my_company}」的商標資料！")
            
            cols_to_show = [c for c in ['商標名稱', '申請人', '申請日期', '類別', '申請案號'] if c in df_real.columns]
            if cols_to_show:
                st.dataframe(my_df[cols_to_show], use_container_width=True, height=500)
            else:
                st.dataframe(my_df, use_container_width=True, height=500)